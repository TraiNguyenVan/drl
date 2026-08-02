import os
import re
import csv
import subprocess
import shutil
import copy
import difflib
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
import docx
import argparse

parser = argparse.ArgumentParser(description="Process and generate DRL scores.")
parser.add_argument("--disable-charts", action="store_true", help="Disable generating DRL charts.")
parser.add_argument("--gdrive-dir", type=str, default=None, help="Path to Google Drive directory for syncing (disabled by default).")
args = parser.parse_args()

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def write_centered_score(cell, text):
    cell.text = text
    if cell.paragraphs:
        cell.paragraphs[0].alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

def write_signature_name(cell, name):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = docx.shared.Pt(0)
    p.paragraph_format.space_after = docx.shared.Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(name)
    run.font.size = docx.shared.Pt(13)
    run.font.name = "Times New Roman"

def format_signature_name(name):
    words = name.split()
    if len(words) <= 3:
        return name
    else:
        line1 = " ".join(words[:2])
        line2 = " ".join(words[2:])
        return f"{line1}\n{line2}"

def fill_info(doc, name, msv, dob_val):
    # Paragraph 4: Name and DOB
    p4 = doc.paragraphs[4]
    p4.text = "" # Clear all text/runs
    
    # Set the custom tab stop on p4 (3.8 inches is 9.65 cm)
    p4.paragraph_format.tab_stops.add_tab_stop(docx.shared.Inches(3.8), docx.enum.text.WD_TAB_ALIGNMENT.LEFT)
    
    r_name_label = p4.add_run("Họ và tên: ")
    r_name_label.font.name = "Times New Roman"
    r_name_label.font.size = docx.shared.Pt(13)
    
    r_name_val = p4.add_run(name)
    r_name_val.font.name = "Times New Roman"
    r_name_val.font.size = docx.shared.Pt(13)
    
    p4.add_run("\t") # Tab to jump to 3.8 inches
    
    r_dob_label = p4.add_run("Ngày sinh: ")
    r_dob_label.font.name = "Times New Roman"
    r_dob_label.font.size = docx.shared.Pt(13)
    
    if dob_val:
        r_dob_val = p4.add_run(dob_val)
        r_dob_val.font.name = "Times New Roman"
        r_dob_val.font.size = docx.shared.Pt(13)
        
    # Paragraph 5: Student ID and Class
    p5 = doc.paragraphs[5]
    p5.text = "" # Clear all text/runs
    
    # Set the custom tab stop on p5
    p5.paragraph_format.tab_stops.add_tab_stop(docx.shared.Inches(3.8), docx.enum.text.WD_TAB_ALIGNMENT.LEFT)
    
    r_msv_label = p5.add_run("Mã số sinh viên: ")
    r_msv_label.font.name = "Times New Roman"
    r_msv_label.font.size = docx.shared.Pt(13)
    
    r_msv_val = p5.add_run(msv)
    r_msv_val.font.name = "Times New Roman"
    r_msv_val.font.size = docx.shared.Pt(13)
    
    p5.add_run("\t") # Tab to jump to 3.8 inches
    
    r_class_label = p5.add_run("Lớp: ")
    r_class_label.font.name = "Times New Roman"
    r_class_label.font.size = docx.shared.Pt(13)
    
    r_class_val = p5.add_run("E25CQCE02-N")
    r_class_val.font.name = "Times New Roman"
    r_class_val.font.size = docx.shared.Pt(13)
            
    # Floating Textbox Student Name replacement (split into 2 lines)
    words = name.split()
    if len(words) <= 3:
        line1 = name
        line2 = ""
    else:
        line1 = " ".join(words[:2])
        line2 = " ".join(words[2:])
        
    p_elms = doc.element.xpath('.//*[local-name()="p"]')
    for p_elm in p_elms:
        p = docx.text.paragraph.Paragraph(p_elm, doc)
        for r in p.runs:
            if "##STUDENT" in r.text:
                r.text = r.text.replace("##STUDENT", line1)
            if "_NAME##" in r.text:
                r.text = r.text.replace("_NAME##", line2)

# Directories
workspace = os.path.dirname(os.path.abspath(__file__))
output_dir = os.path.join(workspace, "generated_students")
os.makedirs(output_dir, exist_ok=True)

# 2. Text Normalization for robust description matching
def normalize_desc(text):
    return re.sub(r'[^a-z0-9àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]', '', text.lower())

def is_max_pts_refined(val):
    val_clean = val.strip().lower()
    if val_clean.endswith("điểm") and len(val_clean) < 15:
        return True
    if val_clean == "100":
        return True
    if val_clean.startswith("-") and len(val_clean) < 6 and any(c.isdigit() for c in val_clean):
        return True
    return False

# 3. Load master template active rows
master_doc_path = os.path.join(workspace, "master.docx")
master_doc = docx.Document(master_doc_path)
master_table = master_doc.tables[1]
master_rows_info = {} # normalized_desc -> {row_idx, max_pts_str}

for r_idx, row in enumerate(master_table.rows):
    unique_cells = []
    seen_tcs = set()
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id not in seen_tcs:
            seen_tcs.add(tc_id)
            unique_cells.append(cell)
            
    max_idx = None
    for idx, cell in enumerate(unique_cells):
        val = cell.text.strip().replace('\n', ' ')
        if is_max_pts_refined(val):
            max_idx = idx
            break
            
    if max_idx is not None and max_idx > 0:
        desc = unique_cells[max_idx - 1].text.strip().replace('\n', ' ')
        norm_desc = normalize_desc(desc)
        master_rows_info[norm_desc] = {
            "row_idx": r_idx,
            "desc": desc,
            "max": unique_cells[max_idx].text.strip()
        }

# Similarity helper for description matching with typos
def get_best_master_match(student_norm_desc, master_keys):
    # Try exact match first
    if student_norm_desc in master_keys:
        return student_norm_desc
    # Fuzzy match
    best_key = None
    best_ratio = 0.0
    for key in master_keys:
        ratio = difflib.SequenceMatcher(None, student_norm_desc, key).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_key = key
    if best_ratio >= 0.8:
        return best_key
    return None

# Helper to parse float scores
def parse_score_val(val):
    if not val:
        return 0.0
    val_clean = val.strip().replace(',', '.')
    val_clean = re.sub(r'[^0-9.-]', '', val_clean)
    try:
        return float(val_clean)
    except ValueError:
        return 0.0

# Helper to classify training point scores into DRL rating tiers
def get_rating(score):
    if score >= 90:
        return "Xuất sắc"
    elif score >= 80:
        return "Tốt"
    elif score >= 65:
        return "Khá"
    elif score >= 50:
        return "Trung bình"
    else:
        return "Yêu"

# 4. Load Excel database (Official Roster List)
db_path = os.path.join(workspace, "HV_Mau 2_Tong hop KQRL cua SV.xlsx")
# Open without data_only=True to preserve formulas and styling on save
wb_db = openpyxl.load_workbook(db_path)
sheet_db = wb_db.active

students_db = []
for r in range(14, sheet_db.max_row + 1):
    tt = sheet_db.cell(r, 1).value
    last_name = sheet_db.cell(r, 2).value
    first_name = sheet_db.cell(r, 3).value
    msv = sheet_db.cell(r, 4).value
    
    is_valid_tt = False
    if isinstance(tt, (int, float)):
        is_valid_tt = True
    elif isinstance(tt, str):
        try:
            float(tt)
            is_valid_tt = True
        except ValueError:
            pass
            
    if is_valid_tt and last_name and first_name and msv:
        msv_str = str(msv).strip()
        if msv_str.startswith("N25"):
            full_name = f"{str(last_name).strip()} {str(first_name).strip()}"
            full_name = " ".join(full_name.split())
            
            # Initialize empty defaults and keep row index for saving updates later
            students_db.append({
                "tt": int(float(tt)),
                "last_name": str(last_name).strip(),
                "first_name": str(first_name).strip(),
                "name": full_name,
                "msv": msv_str,
                "excel_tc": [0.0] * 5,
                "excel_total": 0.0,
                "excel_rating": "",
                "excel_notes": "",
                "row_idx": r
            })

# 4b. Load new vertical format ai_studio_code.csv
csv_path = os.path.join(workspace, "ai_studio_code.csv")
csv_scores_by_msv = {}
csv_dob_by_msv = {}
if os.path.exists(csv_path):
    with open(csv_path, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            msv = row["student_id"].strip()
            crit_id = row["criterion_id"].strip()
            if msv not in csv_scores_by_msv:
                csv_scores_by_msv[msv] = {}
            csv_scores_by_msv[msv][crit_id] = {
                "sv": row["student_score"].strip() if row.get("student_score") is not None else "",
                "lop": row["class_score"].strip() if row.get("class_score") is not None else "",
                "cvht": row["advisor_score"].strip() if row.get("advisor_score") is not None else "",
                "note": row["note"].strip() if row.get("note") is not None else ""
            }
            if row.get("dob"):
                csv_dob_by_msv[msv] = row["dob"].strip()

# Subsection Row Mappings in master.docx Table 1
subsection_mapping = {
    "1.1": [4],
    "1.2": [6, 7, 8, 9, 10, 11],
    "1.3": [12, 14, 15, 16, 17],
    "1.4": [18],
    "1.5": [19],
    "2.1": [22, 24, 25],
    "2.2": [26, 27],
    "2.3": [28, 29],
    "3.1": [32],
    "3.2": [33],
    "3.3": [34],
    "3.4": [35],
    "3.5": [36],
    "4.1": [39],
    "4.2": [40],
    "4.3": [41],
    "4.4": [42],
    "4.5": [43],
    "4.6": [44],
    "5.1": [47],
    "5.2": [48],
    "5.3": [50, 51]
}

# Helper to write score from CSV to table cells
def write_csv_score_to_table(table, criterion_id, role_col, score_str):
    if not score_str:
        return
        
    try:
        val = float(score_str)
        if val.is_integer():
            val_str = str(int(val))
        else:
            val_str = str(val)
    except ValueError:
        val_str = score_str
        val = 0.0

    # Map criterion_id to row index and write
    if criterion_id == "1.1":
        write_centered_score(table.rows[4].cells[role_col], val_str)
    elif criterion_id == "1.2":
        if val == 10:
            write_centered_score(table.rows[6].cells[role_col], "10")
        elif val == 8:
            write_centered_score(table.rows[7].cells[role_col], "8")
        elif val == 6:
            write_centered_score(table.rows[8].cells[role_col], "6")
        elif val == 4:
            write_centered_score(table.rows[9].cells[role_col], "4")
        elif val == 0:
            write_centered_score(table.rows[10].cells[role_col], "0")
    elif criterion_id == "1.3":
        write_centered_score(table.rows[12].cells[role_col], "4")
        if val < 4:
            penalty = int(val - 4)
            write_centered_score(table.rows[14].cells[role_col], str(penalty))
    elif criterion_id == "1.4":
        write_centered_score(table.rows[18].cells[role_col], val_str)
    elif criterion_id == "1.5":
        write_centered_score(table.rows[19].cells[role_col], val_str)
    elif criterion_id == "2.1":
        write_centered_score(table.rows[22].cells[role_col], "15")
        if val < 15:
            penalty = int(val - 15)
            write_centered_score(table.rows[24].cells[role_col], str(penalty))
    elif criterion_id == "2.2":
        write_centered_score(table.rows[26].cells[role_col], "5")
        if val < 5:
            penalty = int(val - 5)
            write_centered_score(table.rows[27].cells[role_col], str(penalty))
    elif criterion_id == "2.3":
        if val > 0:
            write_centered_score(table.rows[28].cells[role_col], val_str)
        else:
            write_centered_score(table.rows[28].cells[role_col], "0")
    elif criterion_id == "3.1":
        write_centered_score(table.rows[32].cells[role_col], val_str)
    elif criterion_id == "3.2":
        write_centered_score(table.rows[33].cells[role_col], val_str)
    elif criterion_id == "3.3":
        write_centered_score(table.rows[34].cells[role_col], val_str)
    elif criterion_id == "3.4":
        write_centered_score(table.rows[35].cells[role_col], val_str)
    elif criterion_id == "3.5":
        write_centered_score(table.rows[36].cells[role_col], val_str)
    elif criterion_id == "4.1":
        write_centered_score(table.rows[39].cells[role_col], val_str)
    elif criterion_id == "4.2":
        write_centered_score(table.rows[40].cells[role_col], val_str)
    elif criterion_id == "4.3":
        write_centered_score(table.rows[41].cells[role_col], val_str)
    elif criterion_id == "4.4":
        write_centered_score(table.rows[42].cells[role_col], val_str)
    elif criterion_id == "4.5":
        write_centered_score(table.rows[43].cells[role_col], val_str)
    elif criterion_id == "4.6":
        write_centered_score(table.rows[44].cells[role_col], val_str)
    elif criterion_id == "5.1":
        write_centered_score(table.rows[47].cells[role_col], val_str)
    elif criterion_id == "5.2":
        write_centered_score(table.rows[48].cells[role_col], val_str)
    elif criterion_id == "5.3":
        if val > 0:
            write_centered_score(table.rows[50].cells[role_col], val_str)
        else:
            write_centered_score(table.rows[50].cells[role_col], "0")
    elif criterion_id == "TC1":
        write_centered_score(table.rows[20].cells[role_col], val_str)
    elif criterion_id == "TC2":
        write_centered_score(table.rows[30].cells[role_col], val_str)
    elif criterion_id == "TC3":
        write_centered_score(table.rows[37].cells[role_col], val_str)
    elif criterion_id == "TC4":
        write_centered_score(table.rows[45].cells[role_col], val_str)
    elif criterion_id == "TC5":
        write_centered_score(table.rows[52].cells[role_col], val_str)
    elif criterion_id == "TOTAL":
        write_centered_score(table.rows[53].cells[role_col], val_str)

# 5. Process each student and generate docx and collect scores
print("\nProcessing student files...")
processed_students = []

for s in students_db:
    # 1. Retrieve DRL totals, rating, and notes dynamically from the CSV instead of Excel
    student_csv = csv_scores_by_msv.get(s["msv"])
    tc_scores = [0.0] * 5
    total_score = 0.0
    rating = "Yếu"
    notes = ""
    
    if student_csv:
        for idx, tc in enumerate(["TC1", "TC2", "TC3", "TC4", "TC5"]):
            val_str = student_csv.get(tc, {}).get("cvht", "0")
            tc_scores[idx] = parse_score_val(val_str)
            
        total_str = student_csv.get("TOTAL", {}).get("cvht", "0")
        total_score = parse_score_val(total_str)
        rating = get_rating(total_score)
        
        # Pull note from TOTAL row if present
        notes = student_csv.get("TOTAL", {}).get("note", "")
        
    s["excel_tc"] = tc_scores
    s["excel_total"] = total_score
    s["excel_rating"] = rating
    s["excel_notes"] = notes
    
    # 2. Write scores back to the source Excel worksheet to compile the final summary
    r_idx = s["row_idx"]
    for idx, score in enumerate(tc_scores):
        sheet_db.cell(r_idx, 5 + idx).value = score
    sheet_db.cell(r_idx, 10).value = total_score
    sheet_db.cell(r_idx, 11).value = rating
    sheet_db.cell(r_idx, 12).value = notes
    
    student_record = copy.deepcopy(s)
    student_record["dob"] = ""
    for sub in subsection_mapping.keys():
        student_record[f"sub_{sub}"] = 0.0
        
    # Retrieve DOB from CSV
    dob_val = csv_dob_by_msv.get(s["msv"], "")
    student_record["dob"] = dob_val
    
    # Check if student is present (has scores in the CSV)
    if not student_csv:
        print(f"  TT={s['tt']}: {s['name']} ({s['msv']}) -> MISSING CSV DATA (Absent)")
    
    # Generate clean word document copy
    dest_doc_path = os.path.join(output_dir, f"{s['name']}_{s['msv']}.docx")
    shutil.copy2(master_doc_path, dest_doc_path)
    doc_out = docx.Document(dest_doc_path)
    
    # Fill personal info paragraphs and student signature name textbox
    fill_info(doc_out, s['name'], s['msv'], dob_val)
    
    # Set all active scores columns in grading table (tables[1]) to empty
    table_out = doc_out.tables[1]
    for r_idx in range(len(table_out.rows)):
        unique_cells = []
        seen_tcs = set()
        for cell in table_out.rows[r_idx].cells:
            tc_id = id(cell._tc)
            if tc_id not in seen_tcs:
                seen_tcs.add(tc_id)
                unique_cells.append(cell)
        max_idx = None
        for idx, cell in enumerate(unique_cells):
            val = cell.text.strip()
            if is_max_pts_refined(val):
                max_idx = idx
                break
        if max_idx is not None and max_idx > 0:
            if max_idx + 1 < len(unique_cells): unique_cells[max_idx + 1].text = ""
            if max_idx + 2 < len(unique_cells): unique_cells[max_idx + 2].text = ""
            if max_idx + 3 < len(unique_cells): unique_cells[max_idx + 3].text = ""
            
    # Write scores from CSV
    if student_csv:
        for crit_id, roles in student_csv.items():
            write_csv_score_to_table(table_out, crit_id, 7, roles["sv"])
            write_csv_score_to_table(table_out, crit_id, 8, roles["lop"])
            write_csv_score_to_table(table_out, crit_id, 11, roles["cvht"])
            
    # Sum up subsection scores from CSV (using Advisor score as final, class/student fallback)
    for sub in subsection_mapping.keys():
        sub_sum = 0.0
        if student_csv and sub in student_csv:
            s_dict = student_csv[sub]
            final_val = s_dict["cvht"] if s_dict["cvht"] else (s_dict["lop"] if s_dict["lop"] else s_dict["sv"])
            sub_sum = parse_score_val(final_val)
        student_record[f"sub_{sub}"] = sub_sum
        
    doc_out.save(dest_doc_path)
    processed_students.append(student_record)

print(f"Successfully generated {len(processed_students)} student Word files.")

# Save the updated final summarized Excel roster
print("Saving final summarized Excel roster...")
wb_db.save(db_path)
print("Roster Excel file saved successfully!")

if not args.disable_charts:
    # 6. Generate DRL statistics charts
    print("\nGenerating DRL statistics charts...")
    from render_charts import generate_all_charts
    chart_base64_images = generate_all_charts(processed_students, workspace)
    
    # 6b. Generate interactive DRL HTML dashboard
    print("Generating DRL interactive HTML dashboard...")
    from render_html import generate_html_dashboard
    generate_html_dashboard(processed_students, workspace, chart_base64_images)

# 7. Cleanup temp files (no-op as temp directory is not used anymore)

# 8. Sync generated files to Google Drive if mounted
if args.gdrive_dir:
    gdrive_dir = args.gdrive_dir
    if os.path.exists(gdrive_dir):
        print(f"\nSyncing generated student files to Google Drive at '{gdrive_dir}'...")
        try:
            # Clear existing files in the GDrive directory to prevent duplicates/obsolete entries
            for filename in os.listdir(gdrive_dir):
                file_path = os.path.join(gdrive_dir, filename)
                if os.path.isfile(file_path):
                    os.remove(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            
            # Copy newly generated files
            for filename in os.listdir(output_dir):
                src_file = os.path.join(output_dir, filename)
                dest_file = os.path.join(gdrive_dir, filename)
                if os.path.isfile(src_file):
                    shutil.copy2(src_file, dest_file)
            print("Sync to Google Drive completed successfully.")
        except Exception as e:
            print(f"Warning: Failed to sync with Google Drive: {e}")
    else:
        print(f"\nWarning: Google Drive sync directory at '{gdrive_dir}' not found. Skipping sync.")

print("Pipeline run finished successfully!")
